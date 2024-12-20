import json
from docx import Document
import os
# Function to create a summary table containing all models and their metrics
def create_summary_table(data, file_path_summary):
    document = Document()
    document.add_heading('Model Performance Metrics', level=1)

    num_models = len(data)
    num_metrics = len(next(iter(data[0].values())))

    # Create a table with rows for metrics and columns for models
    table = document.add_table(rows=num_metrics + 1, cols=num_models + 1)
    table.style = 'Table Grid'

    # Set header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metrics'
    for i, model_data in enumerate(data):
        model_name = list(model_data.keys())[0]
        hdr_cells[i + 1].text = model_name

    # Populate the table rows
    for row_idx, metric_name in enumerate(next(iter(data[0].values())).keys(), start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = metric_name
        for col_idx, model_data in enumerate(data):
            row_cells[col_idx + 1].text = str(model_data[list(model_data.keys())[0]][metric_name])

    document.save(file_path_summary)


# Function to create detailed tables for each model with its metrics
def create_detailed_tables(data, file_path_detailed):
    document = Document()
    for model_data in data:
        for model, metrics in model_data.items():
            document.add_heading(model, level=2)

            # Create a table with rows for each metric
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'

            # Populate the table with the metrics
            for metric, value in metrics.items():
                row_cells = table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = str(value)

    document.save(file_path_detailed)


# Main script to load JSON and generate documents
if __name__ == "__main__":
    # Path to the JSON file
    metrics_dir = os.path.abspath(os.path.join(os.getcwd(), "results/metrics"))
    json_file_path = os.path.join(metrics_dir, "metrics.json")


    docs_directory = os.path.abspath(os.path.join(os.getcwd(), "results/Performance_DOCS"))
    summary_file_path = os.path.join(docs_directory, "model_summary.docx")
    detailed_file_path = os.path.join(docs_directory, "model_details.docx")

    # Load the JSON file
    with open(json_file_path, "r") as file:
        model_metrics = json.load(file)

    # Ensure the data is in the expected format (list of dicts)
    if isinstance(model_metrics, dict):
        model_metrics = [model_metrics]

    # Generate the Word documents
    create_summary_table(model_metrics, summary_file_path)
    create_detailed_tables(model_metrics, detailed_file_path)

    print(f"Summary saved to: {summary_file_path}")
    print(f"Details saved to: {detailed_file_path}")
